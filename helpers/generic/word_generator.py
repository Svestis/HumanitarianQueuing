"""
word_generator.py

Creates/appends plots (images), tables and headers to a word document for summary analysis

Last Updated: 2025_04_06
"""
from typing import Optional
import docx.table
import pandas as pd
from docx import Document
from docx.shared import Inches
import os

def append_images(image_paths: list[str], captions: list[str], heading: Optional[str] = None,
                  doc_path: Optional[str] ="./resources/initial_data.docx") -> None:
    """
    Write images and captions to a word file
    :param image_paths: List of image file paths
    :type image_paths: list[str]
    :param captions: List of captions corresponding to images
    :type captions: list[str]
    :param heading: Optional heading above the images
    :type heading: Optional[str]
    :param doc_path: Path to the word file
    :type doc_path: Optional[str]
    :return: None
    """
    doc: Document = Document(doc_path) if os.path.exists(doc_path) else Document()

    if heading:
        doc.add_heading(heading, level=1)

    assert len(image_paths) == len(captions), ("Num of images and num of captions is not the same. Ensure that "
                                               "every image has a corresponding caption")

    for (path, text) in zip(image_paths, captions):
        if os.path.exists(path):
            doc.add_paragraph()
            doc.add_picture(path, width=Inches(5.5))
            doc.add_paragraph(text, style='Caption')
        else:
            continue

    os.makedirs(os.path.dirname(doc_path), exist_ok=True)
    doc.save(doc_path)


def append_table_from_df(df: pd.DataFrame, heading: Optional[str] = None,
                         doc_path: Optional[str] = "./resources/initial_data.docx") -> None:
    """
    Append a table from df to the word file
    :param df: A dataframe to write as table
    :type df: pd.Dataframe
    :param heading: Heading above the table
    :type heading: Optional[str]
    :param doc_path: Path to the word file
    :type doc_path: Optional[str]
    :return: None
    """
    doc: Document = Document(doc_path) if os.path.exists(doc_path) else Document()

    if heading:
        doc.add_heading(heading, level=1)

    # Create table with headers
    table: docx.table.Table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Light List'

    # Header cells
    hdr_cells: tuple = table.rows[0].cells

    # Fill header cells from df
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)

    # Add each row from the df to the table
    for _, row in df.iterrows():

        # Add a new row and get reference to the cells
        row_cells: tuple = table.add_row().cells

        # Fill rows
        for i, value in enumerate(row):

            # Content formatting
            if isinstance(value, float):
                row_cells[i].text = str(int(value)) if value.is_integer() else str(round(value, 2))
            else:
                row_cells[i].text = str(value)

    os.makedirs(os.path.dirname(doc_path), exist_ok=True)
    doc.save(doc_path)