"""
test_word_generator.py

Validating word_generator.py functionality

Last Updated: 2025_04_06
"""
import pandas as pd
from pathlib import Path
import pytest
from docx import Document
from helpers.generic.word_generator import append_images, append_table_from_df
from PIL import Image

def test_append_table_creates_doc(tmp_path: Path) -> None:
    """
    Testing append table function
    :param tmp_path: The path of the temp folder
    :type tmp_path: Path
    :return: None
    """
    doc_path: Path = tmp_path / "test_doc.docx"
    df = pd.DataFrame({
        "Column1": [1, 2],
        "Column2": ["A", "B"]
    })

    append_table_from_df(df, heading="Test Table", doc_path=str(doc_path))

    assert doc_path.exists()

    doc: Document = Document(str(doc_path))

    assert any("Test Table" in p.text for p in doc.paragraphs)

def test_append_images_creates_doc(tmp_path: Path) -> None:
    """
    Testing appending images function
    :param tmp_path: The path of the temp folder
    :type tmp_path: Path
    :return: None
    """
    # Create dummy image
    img_path: Path = tmp_path / "test_img.png"
    Image.new("RGB", (100, 100), color="red").save(img_path)

    doc_path: Path = tmp_path / "test_doc.docx"
    append_images([str(img_path)], ["Test Image"], heading="Test Images", doc_path=str(doc_path))

    assert doc_path.exists()

    doc = Document(str(doc_path))

    assert any("Test Images" in p.text for p in doc.paragraphs)

def test_append_table_and_images(tmp_path: Path) -> None:
    """
    Running image and table functions sequentially
    :param tmp_path: The path of the temp folder
    :type tmp_path: Path
    :return: None
    """
    df: pd.DataFrame = pd.DataFrame({"A": [1, 2], "B": [3, 4]})

    img_path: Path = tmp_path / "img.png"
    Image.new("RGB", (100, 100)).save(img_path)

    doc_path: Path = tmp_path / "combo.docx"

    append_table_from_df(df, heading="Summary", doc_path=str(doc_path))
    append_images([str(img_path)], ["Sample"], heading="Charts", doc_path=str(doc_path))

    assert doc_path.exists()

    doc = Document(str(doc_path))

    assert any("Summary" in p.text for p in doc.paragraphs)
    assert any("Charts" in p.text for p in doc.paragraphs)

def test_append_images_mismatched_lengths(tmp_path: Path) -> None:
    """
    Testing mismatched number of images + captions
    :param tmp_path: The path to the dir
    :type tmp_path: Path
    :return: None
    """
    img_path = tmp_path / "test_img.png"
    Image.new("RGB", (100, 100)).save(img_path)

    doc_path = tmp_path / "test_mismatch.docx"

    with pytest.raises(AssertionError, match="Num of images and num of captions is not the same"):
        append_images([str(img_path)], [], heading="Mismatch Test", doc_path=str(doc_path))